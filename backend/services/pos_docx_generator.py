"""
POS DOCX Generator — S4 v2
============================
Genera POS fedele al documento aziendale storico.
30 sezioni, header/footer, tabelle con bordi semplici,
stile documentale-formale italiano.

Modalita: bozza_interna | bozza_revisione | finale_stampabile
"""

import io
import logging
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.database import db

logger = logging.getLogger(__name__)

# ═══════════════════════════════════════════════════════════════
#  CONSTANTS & HELPERS
# ═══════════════════════════════════════════════════════════════

BLUE_HEADER = "D6E4F0"
YELLOW_HL = "FFFF00"
LIGHT_GRAY = "F2F2F2"


def _shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell, sz="4", color="000000"):
    """Set thin black borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        borders.append(el)
    tcPr.append(borders)


def _para(doc, text, bold=False, italic=False, size=10, align=None, space_after=6, highlight=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if highlight:
        run.font.highlight_color = 7  # Yellow
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def _heading_section(doc, text, highlight_yellow=True):
    """Section heading — bold, optionally yellow-highlighted, like the real POS."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = True
    if highlight_yellow:
        run.font.highlight_color = 7  # Yellow
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def _sub_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def _bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p


def _simple_table(doc, headers, rows, col_widths=None):
    """Table with thin black borders, no shading (faithful to original POS)."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9)
        _set_cell_borders(cell)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val) if val else "")
            run.font.name = "Arial"
            run.font.size = Pt(9)
            _set_cell_borders(cell)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table


def _kv_table(doc, pairs, col_widths=None):
    """Key-value table (2 columns, label bold, value normal)."""
    table = doc.add_table(rows=len(pairs), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (k, v) in enumerate(pairs):
        c0 = table.rows[i].cells[0]
        c0.text = ""
        r = c0.paragraphs[0].add_run(k)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(9)
        _set_cell_borders(c0)

        c1 = table.rows[i].cells[1]
        c1.text = ""
        r = c1.paragraphs[0].add_run(str(v) if v else "")
        r.font.name = "Arial"
        r.font.size = Pt(9)
        _set_cell_borders(c1)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def _sogg(soggetti, ruolo):
    for s in soggetti:
        if s.get("ruolo") == ruolo:
            return s
    return {}


def _placeholder(mode, field_name=""):
    if mode == "bozza_interna":
        return f"[DA COMPLETARE: {field_name}]"
    elif mode == "bozza_revisione":
        return "Da completare prima dell'emissione finale"
    return ""


def _fmt_date(d):
    if not d:
        return ""
    if isinstance(d, str) and len(d) >= 10:
        return d[:10]
    return str(d)


def _add_header_footer(doc, company_name):
    """Add header with company name + POS title + page number (light blue bg)."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.clear()

        # Left: company name
        run_left = p.add_run(company_name)
        run_left.font.name = "Arial"
        run_left.font.size = Pt(8)
        run_left.bold = True

        # Tab + right: POS title
        run_tab = p.add_run("\t\t")
        run_right = p.add_run("Piano Operativo di Sicurezza")
        run_right.font.name = "Arial"
        run_right.font.size = Pt(8)

        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add thin bottom border to header
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:color"), "4472C4")
        bottom.set(qn("w:space"), "1")
        pBdr.append(bottom)
        pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════
#  SECTION GENERATORS (30 sezioni fedeli al POS reale)
# ═══════════════════════════════════════════════════════════════

def _s01_frontespizio(doc, company, dc, cantiere, mode):
    """Sezione 1 — Frontespizio (cover page grande e professionale)."""
    # Spacing top
    for _ in range(3):
        _para(doc, "", size=12, space_after=12)

    # Main title — grande e centrato
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("P.O.S.")
    run.font.name = "Arial"
    run.font.size = Pt(48)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("PIANO OPERATIVO DI SICUREZZA")
    run2.font.name = "Arial"
    run2.font.size = Pt(18)
    run2.bold = True
    run2.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    p2.paragraph_format.space_after = Pt(8)

    # Legal reference
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(
        "Documento ex artt. 17 - 96 del D.Lgs. 81/2008 e s.m.i.\n"
        "redatto in conformita all'allegato XV del D.Lgs. 81/2008 e s.m.i."
    )
    run3.font.name = "Arial"
    run3.font.size = Pt(10)
    run3.italic = True
    p3.paragraph_format.space_after = Pt(24)

    # Separator line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("_" * 60)
    run_line.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    p_line.paragraph_format.space_after = Pt(16)

    # Company info table
    bname = company.get("business_name", _placeholder(mode, "Ragione sociale"))
    address = company.get("address", _placeholder(mode, "Sede"))
    attivita = company.get("attivita_ateco", "Costruzione di strutture metalliche e parti assemblate")

    _kv_table(doc, [
        ("Azienda", bname),
        ("Sede", address),
        ("Attivita", attivita),
    ], col_widths=[4, 13])

    # Cantiere info
    attivita_cantiere = dc.get("attivita_cantiere", _placeholder(mode, "Attivita cantiere"))
    indirizzo_cantiere = dc.get("indirizzo_cantiere", "")
    citta = dc.get("citta_cantiere", "")
    prov = dc.get("provincia_cantiere", "")
    indirizzo_full = f"{indirizzo_cantiere}, {citta} ({prov})" if indirizzo_cantiere else _placeholder(mode, "Indirizzo cantiere")

    _kv_table(doc, [
        ("Cantiere", attivita_cantiere),
        ("Indirizzo cantiere", indirizzo_full),
    ], col_widths=[4, 13])

    # Revision table
    revs = cantiere.get("revisioni", [{"rev": "00", "motivazione": "Emissione", "data": ""}])
    rev_rows = []
    for r in revs:
        rev_rows.append([r.get("rev", "00"), r.get("motivazione", ""), _fmt_date(r.get("data"))])
    # Pad to 4 rows
    while len(rev_rows) < 4:
        rev_rows.append([f"{len(rev_rows):02d}", "", ""])
    _simple_table(doc, ["Rev.", "Motivazione", "Data"], rev_rows, col_widths=[2, 10, 5])

    doc.add_page_break()


def _s03_introduzione(doc):
    """Sezione 3 — Introduzione (testo fisso aziendale)."""
    _heading_section(doc, "Introduzione")
    _para(doc,
        "Il presente Piano Operativo di Sicurezza (POS) e redatto ai sensi dell'art. 96 comma 1 "
        "lettera g) del D.Lgs. 81/2008 e s.m.i., in conformita all'allegato XV del medesimo decreto. "
        "Esso costituisce il documento di valutazione dei rischi specifici dell'impresa esecutrice "
        "relativamente al cantiere in oggetto.")
    _para(doc,
        "Il POS contiene le informazioni relative alla organizzazione del cantiere, alle lavorazioni "
        "previste, ai rischi specifici connessi e alle misure di prevenzione e protezione adottate, "
        "con particolare riferimento ai dispositivi di protezione individuale (DPI).")
    _para(doc,
        "Il presente documento deve essere:\n"
        "- conservato in cantiere e reso disponibile alle autorita competenti;\n"
        "- consultato dal Coordinatore per l'Esecuzione (CSE) ove nominato;\n"
        "- portato a conoscenza dei lavoratori e dei loro rappresentanti;\n"
        "- aggiornato in caso di variazioni significative delle condizioni operative.")


def _s04_documentazione(doc):
    """Sezione 4 — Elenco documentazione da conservare in cantiere."""
    _heading_section(doc, "Elenco Documentazione da conservare in cantiere")
    _para(doc,
        "Si riporta di seguito un elenco indicativo e non esaustivo della documentazione che "
        "deve essere conservata in cantiere a cura dell'Impresa:")
    docs_list = [
        "PATENTE A CREDITI / iscrizione alla camera di commercio, industria, artigianato e agricoltura / "
        "adempimento degli obblighi formativi / possesso di DURC valido / possesso di DVR valido / "
        "possesso della certificazione di regolarita fiscale / designazione RSPP",
        "Piano Operativo di Sicurezza",
        "Modelli UNILAV",
        "Denuncia INAIL inizio attivita e variazioni",
        "Dichiarazione organico medio annuo e tipo di contratto applicato con i dipendenti",
        "Nomina dei coordinatori dell'emergenza ed elenco dei componenti",
        "Registro degli infortuni debitamente vidimato",
        "Nomina Medico Competente e registro visite mediche dipendenti",
        "Copia invio INAIL/ASL dichiarazione conformita impianto messa a terra e protezione scariche atmosferiche",
        "Denuncia apparecchi di sollevamento portata superiore a 200 kg",
        "Libretti apparecchi di sollevamento con portata superiore a 200 kg",
        "Schede verifiche trimestrali funi e catene",
        "Autorizzazione ministeriale ponteggio metallico / disegno esecutivo se > 20 m",
        "Libretto costruttore ponteggio con limiti di carico e modalita impiego",
        "Schede di sicurezza sostanze chimiche utilizzate",
        "Attestati di formazione sicurezza dei lavoratori",
        "Certificati di idoneita sanitaria dei lavoratori",
    ]
    for d in docs_list:
        _bullet(doc, d)


def _s05_presentazione(doc, company):
    """Sezione 5 — Presentazione dell'azienda (testo fisso + dati azienda)."""
    _heading_section(doc, "Presentazione dell'azienda")
    bn = company.get("business_name", "")
    attivita = company.get("attivita_ateco", "costruzione di strutture metalliche")
    _para(doc,
        f"L'impresa {bn} opera nel settore della {attivita.lower()}. "
        "L'azienda ha maturato esperienza nella realizzazione di strutture in acciaio, "
        "dalla progettazione alla fabbricazione, dal montaggio in opera alla manutenzione. "
        "Le lavorazioni vengono eseguite sia in officina che presso cantieri esterni, "
        "nel rispetto delle normative vigenti in materia di sicurezza sul lavoro e di qualita.")


def _s06_anagrafica(doc, company, soggetti, mode):
    """Sezione 6 — Anagrafica aziendale."""
    _heading_section(doc, "Anagrafica aziendale")

    ddl = _sogg(soggetti, "DATORE_LAVORO")
    rspp = _sogg(soggetti, "RSPP")
    mc = _sogg(soggetti, "MEDICO_COMPETENTE")
    preposto = _sogg(soggetti, "PREPOSTO_CANTIERE")
    dt = _sogg(soggetti, "DIRETTORE_TECNICO")

    rows = [
        ("Azienda", company.get("business_name", _placeholder(mode, "Ragione sociale"))),
        ("Sede legale", company.get("address", "")),
        ("Telefono", company.get("phone", "")),
        ("Email / PEC", company.get("email", "")),
        ("P.IVA / C.F.", company.get("partita_iva", "")),
        ("Datore di Lavoro", ddl.get("nome", _placeholder(mode, "Datore di Lavoro"))),
        ("RSPP", rspp.get("nome", _placeholder(mode, "RSPP"))),
        ("Medico Competente", mc.get("nome", _placeholder(mode, "Medico Competente"))),
        ("Direttore Tecnico", dt.get("nome", "")),
        ("Preposto di Cantiere", preposto.get("nome", "")),
    ]
    _kv_table(doc, rows, col_widths=[5, 12])


def _s07_mansionario(doc, soggetti, lavoratori, mode):
    """Sezione 7 — Mansionario."""
    _heading_section(doc, "Mansionario")
    _para(doc,
        "Si riporta di seguito il mansionario dei lavoratori e delle figure professionali "
        "coinvolte nelle attivita di cantiere.")

    # Key roles
    ddl = _sogg(soggetti, "DATORE_LAVORO")
    rspp = _sogg(soggetti, "RSPP")
    preposto = _sogg(soggetti, "PREPOSTO_CANTIERE")

    roles_rows = [
        ["Datore di Lavoro", ddl.get("nome", ""), "Direzione, organizzazione, vigilanza"],
        ["RSPP", rspp.get("nome", ""), "Servizio prevenzione e protezione"],
        ["Preposto di Cantiere", preposto.get("nome", ""), "Sorveglianza operativa, coordinamento"],
    ]
    _simple_table(doc, ["Ruolo", "Nominativo", "Mansione / Compiti"], roles_rows, col_widths=[5, 5, 7])

    if lavoratori:
        _sub_heading(doc, "Lavoratori assegnati al cantiere")
        worker_rows = []
        for lav in lavoratori:
            worker_rows.append([
                lav.get("nominativo", ""),
                lav.get("mansione", ""),
                "Si" if lav.get("addetto_primo_soccorso") else "No",
                "Si" if lav.get("addetto_antincendio") else "No",
            ])
        _simple_table(doc, ["Nominativo", "Mansione", "Primo Soccorso", "Antincendio"],
                       worker_rows, col_widths=[5, 4.5, 3.5, 3.5])


def _s08_dati_cantiere(doc, dc, mode):
    """Sezione 8 — Dati relativi al cantiere."""
    _heading_section(doc, "Dati relativi al cantiere")
    rows = [
        ("Oggetto dei lavori", dc.get("attivita_cantiere", _placeholder(mode, "Attivita cantiere"))),
        ("Committente", dc.get("committente_nome", _placeholder(mode, "Committente"))),
        ("Indirizzo cantiere", dc.get("indirizzo_cantiere", _placeholder(mode, "Indirizzo"))),
        ("Comune / Provincia", f"{dc.get('citta_cantiere', '')} ({dc.get('provincia_cantiere', '')})"),
        ("Data presunta inizio lavori", _fmt_date(dc.get("data_inizio_lavori")) or _placeholder(mode, "Data inizio")),
        ("Data presunta fine lavori", _fmt_date(dc.get("data_fine_prevista")) or _placeholder(mode, "Data fine")),
        ("Durata presunta", dc.get("durata_prevista", "")),
    ]
    _kv_table(doc, rows, col_widths=[6, 11])


def _s09_soggetti(doc, soggetti, mode):
    """Sezione 9 — Soggetti di riferimento."""
    _heading_section(doc, "Soggetti di riferimento")
    mapping = [
        ("COMMITTENTE", "Committente"),
        ("REFERENTE_COMMITTENTE", "Referente Committente"),
        ("RESPONSABILE_LAVORI", "Responsabile dei Lavori"),
        ("DIRETTORE_LAVORI", "Direttore dei Lavori"),
        ("CSP", "Coordinatore Sicurezza in fase di Progettazione"),
        ("CSE", "Coordinatore Sicurezza in fase di Esecuzione"),
        ("PROGETTISTA", "Progettista"),
        ("STRUTTURISTA", "Strutturista"),
        ("COLLAUDATORE", "Collaudatore"),
    ]
    rows = []
    for ruolo, label in mapping:
        s = _sogg(soggetti, ruolo)
        nome = s.get("nome", "")
        if nome:
            rows.append([label, nome, s.get("telefono", ""), s.get("email", "")])

    if rows:
        _simple_table(doc, ["Ruolo", "Nominativo", "Telefono", "Email"], rows, col_widths=[6, 5, 3, 3])
    else:
        if mode == "finale_stampabile":
            _para(doc, "Non risultano soggetti di riferimento esterni designati per il presente cantiere.")
        else:
            _para(doc, _placeholder(mode, "Soggetti di riferimento"))


def _s10_turni(doc, turni, mode):
    """Sezione 10 — Turni di lavoro."""
    _heading_section(doc, "Turni di lavoro")
    matt = turni.get("mattina", "08:00 - 12:00") if turni else "08:00 - 12:00"
    pom = turni.get("pomeriggio", "13:00 - 17:00") if turni else "13:00 - 17:00"
    _simple_table(doc, ["Turno", "Orario"], [
        ["Mattina", matt],
        ["Pomeriggio", pom],
    ], col_widths=[6, 11])


def _s11_subappalto(doc, subappalti, mode):
    """Sezione 11 — Lavorazioni in subappalto."""
    _heading_section(doc, "Indicazione sulla natura delle lavorazioni da eseguire in subappalto")
    if not subappalti:
        _para(doc, "Non risultano previste lavorazioni in subappalto per il presente cantiere.")
    else:
        rows = [[s.get("lavorazione", ""), s.get("impresa", ""), s.get("durata_prevista", "")] for s in subappalti]
        _simple_table(doc, ["Lavorazione", "Impresa / Lav. Autonomo", "Durata prevista"], rows, col_widths=[6, 6, 5])


def _s12_misure_prevenzione(doc, misure_calc, dpi_lib_map):
    """Sezione 12 — Principali misure di prevenzione (ibrida: fisso + dinamico)."""
    _heading_section(doc, "Principali misure di prevenzione")
    _para(doc,
        "Si riportano di seguito le principali misure di prevenzione e protezione adottate "
        "dall'impresa per la tutela della sicurezza e salute dei lavoratori, conformi a quanto "
        "prescritto dal D.Lgs. 81/2008 e s.m.i.")

    # Fixed measures
    fixed_measures = [
        "Obblighi generali dei lavoratori (Art. 20 D.Lgs. 81/2008)",
        "Prevenzione rischio investimento e caduta materiali dall'alto",
        "Prevenzione scivolamenti, cadute a livello",
        "Protezione dall'esposizione al rumore (Titolo VIII Capo II)",
        "Protezione da tagli, punture, abrasioni",
        "Prevenzione cesoiamento e schiacciamento",
        "Misure per lavori in elevazione e caduta dall'alto",
        "Uso corretto di scale portatili (Art. 113)",
        "Precauzioni per uso vernici e solventi",
        "Movimentazione manuale dei carichi (Titolo VI)",
        "Protezione da polveri e agenti chimici",
        "Sorveglianza sanitaria (Art. 41)",
    ]
    for m in fixed_measures:
        _bullet(doc, m)

    # Dynamic measures from rischi
    dynamic = []
    for m in misure_calc:
        codice = m.get("codice", "")
        lib = dpi_lib_map.get(codice, {})
        if lib.get("tipo") == "misura" and lib.get("nome"):
            dynamic.append(lib["nome"])

    if dynamic:
        _sub_heading(doc, "Misure specifiche attivate dall'analisi dei rischi")
        for m in sorted(set(dynamic)):
            _bullet(doc, m)


def _s13_formazione(doc):
    """Sezione 13 — Attivita formativa."""
    _heading_section(doc, "Attivita Formativa")
    _para(doc,
        "L'impresa garantisce che tutti i lavoratori impiegati nel cantiere hanno ricevuto "
        "adeguata formazione, informazione e addestramento in materia di sicurezza sul lavoro, "
        "ai sensi degli artt. 36 e 37 del D.Lgs. 81/2008 e degli Accordi Stato-Regioni vigenti.")
    _para(doc,
        "In particolare, la formazione comprende:\n"
        "- Formazione generale (4 ore) e specifica (rischio alto 12 ore)\n"
        "- Aggiornamento quinquennale (6 ore)\n"
        "- Formazione uso attrezzature specifiche (PLE, gru, carrelli)\n"
        "- Addestramento all'uso dei DPI di terza categoria\n"
        "- Formazione primo soccorso e antincendio per gli addetti designati\n"
        "- Formazione preposti (8 ore + aggiornamento biennale)")


def _s14_sorveglianza(doc, soggetti, mode):
    """Sezione 14 — Sorveglianza sanitaria."""
    _heading_section(doc, "Sorveglianza Sanitaria")
    mc = _sogg(soggetti, "MEDICO_COMPETENTE")
    mc_nome = mc.get("nome", _placeholder(mode, "Medico Competente"))
    _para(doc,
        f"La sorveglianza sanitaria e effettuata dal Medico Competente, Dott. {mc_nome}, "
        "che provvede alle visite mediche preventive e periodiche dei lavoratori esposti a rischi "
        "specifici, ai sensi dell'art. 41 del D.Lgs. 81/2008.")
    _para(doc,
        "I giudizi di idoneita alla mansione specifica sono conservati in azienda e "
        "i relativi certificati sono disponibili in cantiere.")


def _s15_programma_sanitario(doc):
    """Sezione 15 — Programma sanitario."""
    _heading_section(doc, "Programma sanitario")
    _para(doc,
        "Il programma di sorveglianza sanitaria prevede accertamenti sanitari preventivi e "
        "periodici in relazione ai rischi specifici a cui i lavoratori sono esposti, con "
        "particolare riferimento a:\n"
        "- esposizione al rumore\n"
        "- esposizione a vibrazioni\n"
        "- movimentazione manuale dei carichi\n"
        "- esposizione ad agenti chimici\n"
        "- lavori in quota\n"
        "- idoneita alla mansione specifica")


def _s16_dpi(doc, dpi_calc, dpi_lib_map):
    """Sezione 16 — DPI."""
    _heading_section(doc, "Dispositivi di protezione individuale (D.P.I.)")

    _sub_heading(doc, "Misure di prevenzione e istruzioni per gli addetti")
    _para(doc,
        "I dispositivi di protezione individuale (DPI) devono essere impiegati quando i rischi "
        "non possono essere evitati o sufficientemente ridotti da misure tecniche di prevenzione, "
        "da mezzi di protezione collettiva, da misure, metodi o procedimenti di riorganizzazione del lavoro "
        "(Art. 75 D.Lgs. 81/2008).")

    # DPI base sections (fixed text blocks like the original POS)
    dpi_sections = [
        ("Casco", "Il casco o elmetto deve essere utilizzato in tutte le attivita con rischio di caduta "
         "di materiali dall'alto o di urto con elementi fissi. Deve essere conforme alla norma UNI EN 397."),
        ("Guanti", "I guanti devono essere adeguati al rischio specifico: guanti in pelle per lavorazioni "
         "meccaniche, guanti dielettrici per rischio elettrico, guanti resistenti al calore per saldatura."),
        ("Calzature di sicurezza", "Le calzature di sicurezza devono essere dotate di puntale resistente "
         "allo schiacciamento e suola antiperforazione, conformi alla norma UNI EN ISO 20345."),
        ("Cuffie e tappi auricolari", "I dispositivi di protezione dell'udito devono essere utilizzati "
         "quando il livello di esposizione al rumore supera i valori superiori di azione (85 dB(A))."),
        ("Maschere antipolvere", "Le maschere e i filtri facciali devono essere adeguati al tipo e alla "
         "concentrazione degli inquinanti presenti (polveri, fumi di saldatura, vapori)."),
        ("Occhiali di sicurezza e schermi", "Gli occhiali e gli schermi devono essere utilizzati per "
         "proteggere gli occhi da schegge, radiazioni UV/IR (saldatura), spruzzi chimici."),
        ("Cinture di sicurezza e dispositivi anticaduta", "I sistemi di arresto caduta devono essere "
         "utilizzati per lavori in quota superiori a 2 m quando non sia possibile predisporre protezioni "
         "collettive. Conformi alla norma UNI EN 361 / UNI EN 355."),
    ]
    for title, text in dpi_sections:
        _sub_heading(doc, title)
        _para(doc, text)

    # Dynamic DPI table from rischi
    dpi_from_lib = []
    for d in dpi_calc:
        codice = d.get("codice", "")
        lib = dpi_lib_map.get(codice, {})
        if lib.get("tipo") == "dpi" and lib.get("nome"):
            dpi_from_lib.append([
                lib.get("nome", codice),
                lib.get("norma_riferimento", ""),
                ", ".join(d.get("da_rischi", [])[:3]),
            ])
    if dpi_from_lib:
        _sub_heading(doc, "Scelta del dispositivo in funzione dell'attivita lavorativa")
        _simple_table(doc, ["D.P.I.", "Norma UNI EN", "Rischi collegati"],
                       dpi_from_lib, col_widths=[6, 4, 7])


def _s17_segnaletica(doc):
    """Sezione 17 — Segnaletica di sicurezza."""
    _heading_section(doc, "Segnaletica di sicurezza")
    _para(doc,
        "La segnaletica di sicurezza installata nel cantiere e conforme al Titolo V del "
        "D.Lgs. 81/2008 e alla normativa UNI EN ISO 7010. Essa comprende:")
    items = [
        "Segnali di divieto (accesso non autorizzato, divieto fumare)",
        "Segnali di avvertimento (pericolo caduta, rischio elettrico, carichi sospesi)",
        "Segnali di obbligo (uso DPI, percorsi obbligati)",
        "Segnali di salvataggio (uscite emergenza, punto di raccolta, primo soccorso)",
        "Segnaletica di cantiere (delimitazione aree, segnalazione scavi)",
    ]
    for item in items:
        _bullet(doc, item)


def _s18_macchine(doc, macchine, mode):
    """Sezione 18 — Macchine, attrezzature, impianti."""
    _heading_section(doc, "Macchine - Attrezzature - Impianti")
    if not macchine:
        if mode == "finale_stampabile":
            _para(doc, "Le macchine e attrezzature specifiche saranno definite in fase operativa.")
        else:
            _para(doc, _placeholder(mode, "Elenco macchine e attrezzature"))
    else:
        rows = [[m.get("nome", ""), "Si" if m.get("marcata_ce") else "No",
                 "Si" if m.get("verifiche_periodiche") else "No"] for m in macchine]
        _simple_table(doc, ["Macchina / Attrezzatura", "Marcata CE", "Verifiche periodiche"],
                       rows, col_widths=[7, 4.5, 5.5])


def _s19_21_sostanze_stoccaggio(doc):
    """Sezioni 19-22 — Sostanze chimiche, Agenti biologici, Stoccaggio, Servizi igienici."""
    _heading_section(doc, "Elenco delle sostanze utilizzate - Prodotti chimici")
    _para(doc,
        "Le schede di sicurezza (SDS) dei prodotti chimici utilizzati sono conservate in cantiere "
        "e consultabili dai lavoratori. I principali prodotti utilizzati nelle lavorazioni di "
        "carpenteria metallica comprendono: vernici, solventi, gas di saldatura, olii lubrificanti.")

    _heading_section(doc, "Esposizioni ad agenti biologici")
    _para(doc, "Non risultano previste esposizioni ad agenti biologici rilevanti per le attivita del presente cantiere, "
               "fatte salve le misure igieniche generali previste dal protocollo aziendale.")

    _heading_section(doc, "Stoccaggio materiali e/o rifiuti")
    _para(doc,
        "I materiali vengono stoccati in aree dedicate, delimitate e segnalate. "
        "I rifiuti derivanti dalle lavorazioni (sfridi metallici, imballaggi, residui chimici) "
        "vengono gestiti secondo la normativa vigente in materia di rifiuti (D.Lgs. 152/2006).")

    _heading_section(doc, "Servizi Igienico - Assistenziali")
    _para(doc,
        "Per il cantiere sono previsti servizi igienico-assistenziali adeguati, comprensivi di "
        "acqua potabile, servizi igienici, spogliatoi e locali per il ricovero. "
        "Ove non disponibili in loco, si provvedera con strutture mobili o con l'utilizzo "
        "delle strutture messe a disposizione dal committente.")


def _s23_valutazione_rischi(doc):
    """Sezione 23 — Valutazione rischi (intro + criteri)."""
    _heading_section(doc, "Relazione sulla valutazione dei rischi per la sicurezza e la salute "
                          "durante il lavoro e relativi criteri adottati")

    _sub_heading(doc, "Obiettivo della valutazione")
    _para(doc,
        "La valutazione dei rischi e stata condotta secondo quanto previsto dagli artt. 28 e 29 "
        "del D.Lgs. 81/2008, con l'obiettivo di individuare i pericoli presenti nell'ambiente "
        "di lavoro e stimare il rischio residuo dopo l'applicazione delle misure di prevenzione.")

    _sub_heading(doc, "Individuazione dei rischi di esposizione - Criteri adottati")
    _para(doc,
        "Il metodo adottato si basa sulla formula: R = P x D\n"
        "dove P = Probabilita di accadimento (scala 1-4), D = Entita del Danno (scala 1-4).\n\n"
        "Classificazione del rischio:\n"
        "- R = 1-2: Rischio basso\n"
        "- R = 3-4: Rischio medio\n"
        "- R = 6-8: Rischio alto\n"
        "- R = 9-16: Rischio molto alto")


def _s24_soggetti_esposti(doc, lavoratori, soggetti, mode):
    """Sezione 24 — Individuazione soggetti esposti."""
    _heading_section(doc, "Individuazione dei Soggetti Esposti")
    if lavoratori:
        _para(doc, "I soggetti esposti ai rischi individuati nel presente cantiere sono:")
        for lav in lavoratori:
            _bullet(doc, f"{lav.get('nominativo', '')} - {lav.get('mansione', '')}")
    else:
        preposto = _sogg(soggetti, "PREPOSTO_CANTIERE")
        if preposto.get("nome"):
            _para(doc, f"I soggetti esposti includono il preposto di cantiere ({preposto['nome']}) "
                       "e i lavoratori assegnati alle operazioni di cantiere.")
        else:
            _para(doc, _placeholder(mode, "Soggetti esposti") if mode != "finale_stampabile"
                  else "I soggetti esposti sono i lavoratori assegnati alle operazioni di cantiere.")


def _s25_28_rischi_specifici(doc):
    """Sezioni 25-28 — Rischi specifici (rumore, vibrazioni, chimico, MMC)."""
    _heading_section(doc, "Valutazione del Rischio Rumore")
    _para(doc,
        "La valutazione del rischio rumore e stata effettuata ai sensi del Titolo VIII Capo II "
        "del D.Lgs. 81/2008. Le attivita di carpenteria metallica comportano livelli di esposizione "
        "che possono superare i valori superiori di azione. Sono previsti DPI uditivi adeguati e "
        "sorveglianza sanitaria specifica.")

    _heading_section(doc, "Valutazione del Rischio Vibrazioni")
    _para(doc,
        "La valutazione del rischio vibrazioni e stata effettuata ai sensi del Titolo VIII Capo III "
        "del D.Lgs. 81/2008, con riferimento sia al sistema mano-braccio (HAV) che al corpo intero (WBV). "
        "Le attrezzature vibranti utilizzate sono state censite e i valori di esposizione sono stati "
        "confrontati con i valori di azione e limite.")

    _heading_section(doc, "Valutazione del Rischio Chimico")
    _para(doc,
        "La valutazione del rischio chimico e stata effettuata ai sensi del Titolo IX del D.Lgs. 81/2008. "
        "I principali agenti chimici presenti nelle lavorazioni sono: fumi di saldatura, polveri metalliche, "
        "solventi e vernici. Le schede di sicurezza sono consultabili in cantiere.")

    _heading_section(doc, "Valutazione del Rischio da Movimentazione Manuale dei carichi")
    _para(doc,
        "La valutazione e stata effettuata ai sensi del Titolo VI del D.Lgs. 81/2008 e con "
        "il metodo NIOSH. Le operazioni di movimentazione manuale in cantiere riguardano "
        "prevalentemente il posizionamento di elementi strutturali e materiali di consumo. "
        "Per carichi superiori a 25 kg si utilizzano mezzi di sollevamento meccanici.")


def _s29_schede_rischio(doc, fasi, rischi_lib_map, dpi_lib_map):
    """Sezione 29 — Schede rischio per fase di lavoro (CUORE DEL POS)."""
    _heading_section(doc, "Individuazione dei rischi, delle misure di prevenzione e di protezione, "
                          "dei dispositivi di protezione individuale")

    if not fasi:
        _para(doc, "Nessuna fase lavorativa individuata. La sezione verra compilata a seguito "
                   "dell'analisi delle attivita di cantiere.")
        return

    for idx, fase in enumerate(fasi, 1):
        codice = fase.get("fase_codice", "")
        confidence = fase.get("confidence", "dedotto")
        rischi_att = fase.get("rischi_attivati", [])

        _sub_heading(doc, f"Fase {idx}: {codice}")

        if fase.get("reasoning"):
            _para(doc, f"Descrizione: {fase['reasoning']}", italic=True, size=9, space_after=4)

        if not rischi_att:
            _para(doc, "Nessun rischio specifico attivato per questa fase.")
            continue

        # Risk table
        risk_rows = []
        dpi_for_fase = set()
        misure_for_fase = []

        for ra in rischi_att:
            rc = ra.get("rischio_codice", "")
            lib = rischi_lib_map.get(rc, {})
            vd = lib.get("valutazione_default", {})
            nome = lib.get("nome", rc)
            risk_rows.append([nome, str(vd.get("probabilita", "")), str(vd.get("entita_danno", "")), vd.get("classe", "")])

            for d in lib.get("dpi_ids", []):
                dpi_for_fase.add(d)
            for m in lib.get("misure_ids", []):
                mlib = dpi_lib_map.get(m, {})
                if mlib.get("nome"):
                    misure_for_fase.append(mlib["nome"])

        _simple_table(doc, ["Rischio", "P", "D", "Classe"], risk_rows, col_widths=[7, 2, 2, 6])

        # DPI
        if dpi_for_fase:
            dpi_rows = []
            for d_code in sorted(dpi_for_fase):
                dlib = dpi_lib_map.get(d_code, {})
                if dlib.get("nome"):
                    dpi_rows.append([dlib["nome"], dlib.get("norma_riferimento", "")])
            if dpi_rows:
                _para(doc, "DPI richiesti:", bold=True, size=9, space_after=2)
                _simple_table(doc, ["D.P.I.", "Norma"], dpi_rows, col_widths=[9, 8])

        # Misure
        if misure_for_fase:
            _para(doc, "Misure di prevenzione:", bold=True, size=9, space_after=2)
            for m in sorted(set(misure_for_fase)):
                _bullet(doc, m, size=9)


def _s30_emergenza(doc, numeri_utili, soggetti, mode):
    """Sezione 30 — Gestione emergenza + numeri utili + dichiarazione."""
    _heading_section(doc, "Gestione dell'emergenza")
    _para(doc,
        "In caso di emergenza i lavoratori devono attenersi alle procedure aziendali previste "
        "e contattare immediatamente i servizi di soccorso. Il preposto di cantiere coordina "
        "le operazioni di evacuazione e messa in sicurezza dell'area.")

    _sub_heading(doc, "Mezzi antincendio")
    _para(doc,
        "Il cantiere e dotato di estintori a polvere e/o CO2 posizionati in punti accessibili. "
        "Gli estintori sono sottoposti a verifica semestrale e revisione periodica.")

    _sub_heading(doc, "Pronto soccorso")
    _para(doc,
        "In cantiere e presente una cassetta di primo soccorso conforme al DM 388/2003. "
        "Gli addetti al primo soccorso designati sono formati con corso di 12 ore (Gruppo B).")

    _sub_heading(doc, "Numeri utili")
    default_nums = [
        ["Vigili del Fuoco", "115"],
        ["Pronto Soccorso / Emergenza Sanitaria", "118"],
        ["Carabinieri", "112"],
        ["Polizia di Stato", "113"],
        ["Guardia di Finanza", "117"],
    ]
    nums = [[n.get("servizio", ""), n.get("numero", "")] for n in numeri_utili] if numeri_utili else default_nums
    _simple_table(doc, ["Servizio", "Numero"], nums, col_widths=[11, 6])


def _s_dichiarazione(doc, company, soggetti, data_dich, mode):
    """Dichiarazione finale."""
    doc.add_page_break()
    _heading_section(doc, "Dichiarazione")

    ddl = _sogg(soggetti, "DATORE_LAVORO")
    rspp = _sogg(soggetti, "RSPP")
    ddl_nome = ddl.get("nome", _placeholder(mode, "Datore di Lavoro"))
    bn = company.get("business_name", "")

    _para(doc,
        f"Il sottoscritto, {ddl_nome}, in qualita di datore di lavoro della Impresa {bn}")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DICHIARA")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

    _para(doc,
        "che il procedimento sulla valutazione dei rischi ex art. 17 del D.Lgs. 81/2008 e s.m.i., "
        "e stato attuato in collaborazione con il Servizio di Prevenzione e Protezione dai rischi, "
        "con il Medico Competente previa consultazione del Rappresentante dei lavoratori per la sicurezza.")

    _para(doc, f"Data: {_fmt_date(data_dich) or '___/___/______'}", size=10, space_after=20)

    rspp_nome = rspp.get("nome", _placeholder(mode, "RSPP"))

    sig_rows = [
        ["Il Datore di Lavoro", ddl_nome, "_________________________"],
        ["Il Responsabile del S.P.P.", rspp_nome, "_________________________"],
        ["Il Rappresentante dei\nlavoratori per la sicurezza", "", "_________________________"],
    ]
    _simple_table(doc, ["Ruolo", "Nominativo", "Firma"], sig_rows, col_widths=[6, 5, 6])


# ═══════════════════════════════════════════════════════════════
#  MAIN GENERATOR
# ═══════════════════════════════════════════════════════════════

async def genera_pos_docx(cantiere_id: str, user_id: str, mode: str = "bozza_revisione") -> dict:
    """Generate POS DOCX faithful to the company's real POS document."""

    if mode not in ("bozza_interna", "bozza_revisione", "finale_stampabile"):
        mode = "bozza_revisione"

    # 1. Load cantiere
    cantiere = await db.cantieri_sicurezza.find_one(
        {"cantiere_id": cantiere_id, "user_id": user_id}, {"_id": 0}
    )
    if not cantiere:
        return {"error": "Cantiere non trovato"}

    # 2. Gate check
    gate = cantiere.get("gate_pos_status", {})
    completezza = gate.get("completezza_percentuale", 0)

    if mode == "finale_stampabile":
        blockers = gate.get("blockers", [])
        if blockers:
            return {
                "error": f"Impossibile generare in modalita finale: {len(blockers)} blockers attivi",
                "blockers": blockers,
            }

    # 3. Load data
    company = await db.company_settings.find_one({"user_id": user_id}, {"_id": 0}) or {}

    rischi_lib_map = {}
    async for r in db.lib_rischi_sicurezza.find({"user_id": user_id, "active": True}, {"_id": 0}):
        rischi_lib_map[r["codice"]] = r

    dpi_lib_map = {}
    async for d in db.lib_dpi_misure.find({"user_id": user_id, "active": True}, {"_id": 0}):
        dpi_lib_map[d["codice"]] = d

    dc = cantiere.get("dati_cantiere", {})
    soggetti = cantiere.get("soggetti", [])
    lavoratori = cantiere.get("lavoratori_coinvolti", [])
    fasi = cantiere.get("fasi_lavoro_selezionate", [])
    dpi_calc = cantiere.get("dpi_calcolati", [])
    misure_calc = cantiere.get("misure_calcolate", [])
    macchine = cantiere.get("macchine_attrezzature", [])
    numeri_utili = cantiere.get("numeri_utili", [])
    turni = cantiere.get("turni_lavoro", {})
    subappalti = cantiere.get("subappalti", [])

    # Inject committente name into dati_cantiere for convenience
    comm = _sogg(soggetti, "COMMITTENTE")
    if comm.get("nome") and not dc.get("committente_nome"):
        dc["committente_nome"] = comm["nome"]

    # 4. Build document
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header
    _add_header_footer(doc, company.get("business_name", ""))

    # ── Generate all 30 sections ──
    _s01_frontespizio(doc, company, dc, cantiere, mode)
    # (Indice omesso in generazione automatica — si genera in Word con Ctrl+Shift+F9)
    _s03_introduzione(doc)
    _s04_documentazione(doc)
    _s05_presentazione(doc, company)
    _s06_anagrafica(doc, company, soggetti, mode)
    _s07_mansionario(doc, soggetti, lavoratori, mode)
    _s08_dati_cantiere(doc, dc, mode)
    _s09_soggetti(doc, soggetti, mode)
    _s10_turni(doc, turni, mode)
    _s11_subappalto(doc, subappalti, mode)
    _s12_misure_prevenzione(doc, misure_calc, dpi_lib_map)
    _s13_formazione(doc)
    _s14_sorveglianza(doc, soggetti, mode)
    _s15_programma_sanitario(doc)
    _s16_dpi(doc, dpi_calc, dpi_lib_map)
    _s17_segnaletica(doc)
    _s18_macchine(doc, macchine, mode)
    _s19_21_sostanze_stoccaggio(doc)
    _s23_valutazione_rischi(doc)
    _s24_soggetti_esposti(doc, lavoratori, soggetti, mode)
    _s25_28_rischi_specifici(doc)
    _s29_schede_rischio(doc, fasi, rischi_lib_map, dpi_lib_map)
    _s30_emergenza(doc, numeri_utili, soggetti, mode)
    _s_dichiarazione(doc, company, soggetti, cantiere.get("data_dichiarazione"), mode)

    # 5. Save
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # 6. Metadata
    now = datetime.now(timezone.utc).isoformat()
    gen_meta = {
        "timestamp": now,
        "mode": mode,
        "completezza_al_momento": completezza,
        "n_fasi": len(fasi),
        "n_rischi": sum(len(f.get("rischi_attivati", [])) for f in fasi),
        "n_pagine_stimate": 30 + len(fasi) * 2,
        "versione": (cantiere.get("pos_generazioni", [{}])[-1].get("versione", 0) + 1)
            if cantiere.get("pos_generazioni") else 1,
    }

    await db.cantieri_sicurezza.update_one(
        {"cantiere_id": cantiere_id, "user_id": user_id},
        {
            "$push": {"pos_generazioni": gen_meta},
            "$set": {"ultima_generazione_pos": now, "updated_at": now},
        }
    )

    attivita = dc.get("attivita_cantiere", "cantiere")[:40].replace(" ", "_")
    filename = f"POS_{attivita}_{now[:10]}.docx"

    return {
        "success": True,
        "file_bytes": buffer.getvalue(),
        "filename": filename,
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "generazione": gen_meta,
        "gate_completezza": completezza,
    }
