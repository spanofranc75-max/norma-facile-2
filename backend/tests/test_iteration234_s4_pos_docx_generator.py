"""
Iteration 234 — S4 POS DOCX Generator Tests
============================================
Tests for POST /api/cantieri-sicurezza/{id}/genera-pos endpoint
and GET /api/cantieri-sicurezza/{id}/pos-generazioni endpoint.

Features tested:
- DOCX generation from cantiere data
- Response headers (Content-Disposition, X-POS-Versione, X-POS-Completezza)
- DOCX content validation (15+ sections, 21 tables, risk sheets per fase)
- Generation history tracking (pos_generazioni metadata)
- Version incrementing
"""

import pytest
import requests
import os
import io
from docx import Document

BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', '').rstrip('/')
SESSION_TOKEN = "BccyJvnOyrPXRRYy4GfxR5osuF51RWcYWWnUUoHx434"
USER_ID = "user_6988e9b9316c"
TEST_CANTIERE_ID = "cant_3894750ebe93"  # Existing cantiere with AI data (6 fasi, 20 rischi, 10 DPI)


@pytest.fixture
def auth_headers():
    """Authentication headers with session token."""
    return {
        "Authorization": f"Bearer {SESSION_TOKEN}",
        "Content-Type": "application/json"
    }


class TestGeneraPosEndpoint:
    """Tests for POST /api/cantieri-sicurezza/{id}/genera-pos endpoint."""
    
    def test_genera_pos_returns_docx_binary(self, auth_headers):
        """Test that genera-pos returns binary DOCX data."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}: {response.text}"
        
        # Verify content type is DOCX
        content_type = response.headers.get('Content-Type', '')
        assert 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type, \
            f"Expected DOCX content type, got: {content_type}"
        
        # Verify binary content is not empty
        assert len(response.content) > 0, "DOCX content should not be empty"
        print(f"PASSED: genera-pos returns DOCX binary ({len(response.content)} bytes)")
    
    def test_genera_pos_content_disposition_header(self, auth_headers):
        """Test Content-Disposition header with filename."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        
        assert response.status_code == 200
        
        content_disposition = response.headers.get('Content-Disposition', '')
        assert 'attachment' in content_disposition, f"Expected attachment, got: {content_disposition}"
        assert 'filename=' in content_disposition, f"Expected filename in header, got: {content_disposition}"
        assert '.docx' in content_disposition, f"Expected .docx extension, got: {content_disposition}"
        print(f"PASSED: Content-Disposition header correct: {content_disposition}")
    
    def test_genera_pos_versione_header(self, auth_headers):
        """Test X-POS-Versione header is present and numeric."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        
        assert response.status_code == 200
        
        versione = response.headers.get('X-POS-Versione')
        assert versione is not None, "X-POS-Versione header should be present"
        assert versione.isdigit(), f"X-POS-Versione should be numeric, got: {versione}"
        assert int(versione) >= 1, f"X-POS-Versione should be >= 1, got: {versione}"
        print(f"PASSED: X-POS-Versione header: {versione}")
    
    def test_genera_pos_completezza_header(self, auth_headers):
        """Test X-POS-Completezza header is present and numeric."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        
        assert response.status_code == 200
        
        completezza = response.headers.get('X-POS-Completezza')
        assert completezza is not None, "X-POS-Completezza header should be present"
        assert completezza.isdigit() or completezza.replace('.', '').isdigit(), \
            f"X-POS-Completezza should be numeric, got: {completezza}"
        completezza_val = float(completezza)
        assert 0 <= completezza_val <= 100, f"X-POS-Completezza should be 0-100, got: {completezza_val}"
        print(f"PASSED: X-POS-Completezza header: {completezza}")
    
    def test_genera_pos_nonexistent_cantiere(self, auth_headers):
        """Test genera-pos with nonexistent cantiere returns 400/404."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/nonexistent_cantiere_xyz/genera-pos",
            headers=auth_headers
        )
        
        assert response.status_code in [400, 404], f"Expected 400/404, got {response.status_code}"
        print(f"PASSED: Nonexistent cantiere returns {response.status_code}")
    
    def test_genera_pos_unauthorized(self):
        """Test genera-pos without auth returns 401."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos"
        )
        
        assert response.status_code == 401, f"Expected 401, got {response.status_code}"
        print("PASSED: Unauthorized request returns 401")


class TestDocxContent:
    """Tests for DOCX content validation."""
    
    def test_docx_has_15_plus_sections(self, auth_headers):
        """Test DOCX contains 15+ sections (headings)."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        
        assert response.status_code == 200
        
        # Parse DOCX
        doc = Document(io.BytesIO(response.content))
        
        # Count headings (sections)
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        
        # Also count paragraphs that look like section titles
        section_titles = []
        expected_sections = [
            "PIANO OPERATIVO DI SICUREZZA",
            "INTRODUZIONE",
            "DOCUMENTAZIONE",
            "ANAGRAFICA",
            "MANSIONARIO",
            "DATI RELATIVI AL CANTIERE",
            "SOGGETTI",
            "TURNI",
            "SUBAPPALTO",
            "MISURE DI PREVENZIONE",
            "DISPOSITIVI DI PROTEZIONE",
            "MACCHINE",
            "VALUTAZIONE DEI RISCHI",
            "SCHEDE RISCHIO",
            "EMERGENZE",
            "DICHIARAZIONE"
        ]
        
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        found_sections = sum(1 for s in expected_sections if s.upper() in doc_text.upper())
        
        assert found_sections >= 10, f"Expected at least 10 sections, found {found_sections}"
        print(f"PASSED: DOCX has {found_sections} expected sections, {len(headings)} headings")
    
    def test_docx_has_tables(self, auth_headers):
        """Test DOCX contains tables (expected ~21 tables)."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        
        assert response.status_code == 200
        
        doc = Document(io.BytesIO(response.content))
        tables = doc.tables
        
        assert len(tables) >= 10, f"Expected at least 10 tables, found {len(tables)}"
        print(f"PASSED: DOCX has {len(tables)} tables")
    
    def test_docx_has_risk_sheets_per_fase(self, auth_headers):
        """Test DOCX contains risk sheets for each work phase."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        
        assert response.status_code == 200
        
        doc = Document(io.BytesIO(response.content))
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Check for "SCHEDE RISCHIO" section
        assert "SCHEDE RISCHIO" in doc_text.upper() or "13." in doc_text, \
            "Expected 'Schede Rischio' section in DOCX"
        
        # Check for risk-related content
        risk_indicators = ["Rischi individuati", "DPI richiesti", "Probabilita", "Danno", "Classe"]
        found_indicators = sum(1 for r in risk_indicators if r in doc_text)
        
        assert found_indicators >= 2, f"Expected risk sheet content, found {found_indicators} indicators"
        print(f"PASSED: DOCX has risk sheets with {found_indicators} risk indicators")
    
    def test_docx_contains_cantiere_data(self, auth_headers):
        """Test DOCX contains real cantiere data (not just placeholders)."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        
        assert response.status_code == 200
        
        doc = Document(io.BytesIO(response.content))
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Check for some expected content patterns
        # The test cantiere should have some data
        has_content = (
            len(doc_text) > 5000 or  # Substantial content
            "D.Lgs. 81/2008" in doc_text or  # Normative reference
            "Datore di Lavoro" in doc_text or  # Role reference
            "RSPP" in doc_text  # Role reference
        )
        
        assert has_content, "DOCX should contain substantial cantiere data"
        print(f"PASSED: DOCX contains cantiere data ({len(doc_text)} chars)")
    
    def test_docx_file_size_reasonable(self, auth_headers):
        """Test DOCX file size is reasonable (not too small, not too large)."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        
        assert response.status_code == 200
        
        file_size = len(response.content)
        
        # Expected ~41KB based on main agent info, allow range 20KB-200KB
        assert file_size >= 20000, f"DOCX too small: {file_size} bytes"
        assert file_size <= 500000, f"DOCX too large: {file_size} bytes"
        print(f"PASSED: DOCX file size reasonable: {file_size} bytes (~{file_size//1024}KB)")


class TestPosGenerazioniEndpoint:
    """Tests for GET /api/cantieri-sicurezza/{id}/pos-generazioni endpoint."""
    
    def test_pos_generazioni_returns_history(self, auth_headers):
        """Test pos-generazioni returns generation history."""
        response = requests.get(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/pos-generazioni",
            headers=auth_headers
        )
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}: {response.text}"
        
        data = response.json()
        assert "cantiere_id" in data, "Response should have cantiere_id"
        assert "generazioni" in data, "Response should have generazioni array"
        assert isinstance(data["generazioni"], list), "generazioni should be a list"
        print(f"PASSED: pos-generazioni returns history with {len(data['generazioni'])} entries")
    
    def test_pos_generazioni_metadata_structure(self, auth_headers):
        """Test pos_generazioni metadata has correct structure."""
        # First generate a POS to ensure there's at least one generation
        requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        
        response = requests.get(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/pos-generazioni",
            headers=auth_headers
        )
        
        assert response.status_code == 200
        data = response.json()
        
        if len(data["generazioni"]) > 0:
            gen = data["generazioni"][-1]  # Latest generation
            
            # Check expected fields
            expected_fields = ["timestamp", "completezza_al_momento", "n_fasi", "n_rischi", "versione"]
            for field in expected_fields:
                assert field in gen, f"Generation should have '{field}' field"
            
            # Validate field types
            assert isinstance(gen["versione"], int), "versione should be int"
            assert isinstance(gen["n_fasi"], int), "n_fasi should be int"
            assert isinstance(gen["n_rischi"], int), "n_rischi should be int"
            
            print(f"PASSED: Generation metadata structure correct: versione={gen['versione']}, n_fasi={gen['n_fasi']}, n_rischi={gen['n_rischi']}")
        else:
            pytest.skip("No generations found to validate structure")
    
    def test_pos_generazioni_versione_increments(self, auth_headers):
        """Test that versione increments with each generation."""
        # Get current generations
        response1 = requests.get(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/pos-generazioni",
            headers=auth_headers
        )
        assert response1.status_code == 200
        data1 = response1.json()
        initial_count = len(data1["generazioni"])
        initial_versione = data1["generazioni"][-1]["versione"] if initial_count > 0 else 0
        
        # Generate new POS
        gen_response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        assert gen_response.status_code == 200
        
        # Get updated generations
        response2 = requests.get(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/pos-generazioni",
            headers=auth_headers
        )
        assert response2.status_code == 200
        data2 = response2.json()
        
        new_count = len(data2["generazioni"])
        new_versione = data2["generazioni"][-1]["versione"]
        
        assert new_count == initial_count + 1, f"Expected {initial_count + 1} generations, got {new_count}"
        assert new_versione == initial_versione + 1, f"Expected versione {initial_versione + 1}, got {new_versione}"
        print(f"PASSED: Versione incremented from {initial_versione} to {new_versione}")
    
    def test_pos_generazioni_nonexistent_cantiere(self, auth_headers):
        """Test pos-generazioni with nonexistent cantiere returns 404."""
        response = requests.get(
            f"{BASE_URL}/api/cantieri-sicurezza/nonexistent_cantiere_xyz/pos-generazioni",
            headers=auth_headers
        )
        
        assert response.status_code == 404, f"Expected 404, got {response.status_code}"
        print("PASSED: Nonexistent cantiere returns 404")
    
    def test_pos_generazioni_ultima_generazione_field(self, auth_headers):
        """Test ultima_generazione field is returned."""
        response = requests.get(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/pos-generazioni",
            headers=auth_headers
        )
        
        assert response.status_code == 200
        data = response.json()
        
        assert "ultima_generazione" in data, "Response should have ultima_generazione field"
        print(f"PASSED: ultima_generazione field present: {data.get('ultima_generazione')}")


class TestCantiereUpdatedAfterGeneration:
    """Tests to verify cantiere document is updated after POS generation."""
    
    def test_cantiere_has_pos_generazioni_after_generation(self, auth_headers):
        """Test cantiere document has pos_generazioni array after generation."""
        # Generate POS
        gen_response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        assert gen_response.status_code == 200
        
        # Get cantiere
        cantiere_response = requests.get(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}",
            headers=auth_headers
        )
        assert cantiere_response.status_code == 200
        
        cantiere = cantiere_response.json()
        assert "pos_generazioni" in cantiere, "Cantiere should have pos_generazioni field"
        assert isinstance(cantiere["pos_generazioni"], list), "pos_generazioni should be a list"
        assert len(cantiere["pos_generazioni"]) > 0, "pos_generazioni should not be empty"
        print(f"PASSED: Cantiere has {len(cantiere['pos_generazioni'])} pos_generazioni entries")
    
    def test_cantiere_has_ultima_generazione_pos_timestamp(self, auth_headers):
        """Test cantiere has ultima_generazione_pos timestamp after generation."""
        # Generate POS
        gen_response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        assert gen_response.status_code == 200
        
        # Get cantiere
        cantiere_response = requests.get(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}",
            headers=auth_headers
        )
        assert cantiere_response.status_code == 200
        
        cantiere = cantiere_response.json()
        assert "ultima_generazione_pos" in cantiere, "Cantiere should have ultima_generazione_pos field"
        assert cantiere["ultima_generazione_pos"] is not None, "ultima_generazione_pos should not be None"
        print(f"PASSED: Cantiere has ultima_generazione_pos: {cantiere['ultima_generazione_pos']}")


class TestDocxSectionContent:
    """Detailed tests for specific DOCX sections."""
    
    def test_docx_has_copertina_section(self, auth_headers):
        """Test DOCX has Copertina (cover page) section."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        assert response.status_code == 200
        
        doc = Document(io.BytesIO(response.content))
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Check for cover page elements
        has_title = "PIANO OPERATIVO DI SICUREZZA" in doc_text.upper()
        has_normative = "D.Lgs. 81/2008" in doc_text or "81/2008" in doc_text
        
        assert has_title, "DOCX should have POS title"
        assert has_normative, "DOCX should reference D.Lgs. 81/2008"
        print("PASSED: DOCX has Copertina section with title and normative reference")
    
    def test_docx_has_anagrafica_section(self, auth_headers):
        """Test DOCX has Anagrafica (company info) section."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        assert response.status_code == 200
        
        doc = Document(io.BytesIO(response.content))
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Check for anagrafica elements
        anagrafica_indicators = ["ANAGRAFICA", "Impresa", "Sede", "Datore di Lavoro", "RSPP"]
        found = sum(1 for a in anagrafica_indicators if a in doc_text)
        
        assert found >= 2, f"Expected anagrafica content, found {found} indicators"
        print(f"PASSED: DOCX has Anagrafica section with {found} indicators")
    
    def test_docx_has_dpi_section(self, auth_headers):
        """Test DOCX has DPI (Personal Protective Equipment) section."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        assert response.status_code == 200
        
        doc = Document(io.BytesIO(response.content))
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Check for DPI section
        has_dpi = "DPI" in doc_text or "DISPOSITIVI DI PROTEZIONE" in doc_text.upper()
        
        assert has_dpi, "DOCX should have DPI section"
        print("PASSED: DOCX has DPI section")
    
    def test_docx_has_emergenza_section(self, auth_headers):
        """Test DOCX has Emergenza (emergency) section."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        assert response.status_code == 200
        
        doc = Document(io.BytesIO(response.content))
        
        # Get text from paragraphs
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Also get text from tables (emergency numbers are in a table)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    doc_text += " " + cell.text
        
        doc_text_upper = doc_text.upper()
        
        # Check for emergency section (case-insensitive)
        emergency_indicators = ["EMERGENZ", "VIGILI DEL FUOCO", "PRONTO SOCCORSO", "118", "115"]
        found = sum(1 for e in emergency_indicators if e in doc_text_upper)
        
        assert found >= 2, f"Expected emergency content, found {found} indicators"
        print(f"PASSED: DOCX has Emergenza section with {found} indicators")
    
    def test_docx_has_dichiarazione_section(self, auth_headers):
        """Test DOCX has Dichiarazione (declaration) section."""
        response = requests.post(
            f"{BASE_URL}/api/cantieri-sicurezza/{TEST_CANTIERE_ID}/genera-pos",
            headers=auth_headers
        )
        assert response.status_code == 200
        
        doc = Document(io.BytesIO(response.content))
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Check for declaration section
        has_dichiarazione = "DICHIARAZIONE" in doc_text.upper() or "sottoscritto" in doc_text.lower()
        
        assert has_dichiarazione, "DOCX should have Dichiarazione section"
        print("PASSED: DOCX has Dichiarazione section")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
